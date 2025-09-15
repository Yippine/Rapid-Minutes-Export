"""
Word Template Engine (D1 - Document Processing Layer)
Advanced Word document template processing and manipulation
Implements ICE principle - Intuitive template processing with comprehensive document handling
"""

import logging
import os
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - Word processing will be limited")

from ..config import settings

logger = logging.getLogger(__name__)


@dataclass
class TemplateField:
    """Template field definition"""
    field_name: str
    placeholder: str
    field_type: str = "text"  # text, table, list, image
    required: bool = True
    default_value: Any = None
    format_options: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ProcessingResult:
    """Word processing result"""
    success: bool
    output_path: Optional[str] = None
    processed_fields: List[str] = field(default_factory=list)
    error_message: Optional[str] = None
    warnings: List[str] = field(default_factory=list)
    processing_time: Optional[float] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


class WordEngine:
    """
    Advanced Word document template processor
    Handles template loading, field replacement, formatting, and document generation
    """
    
    def __init__(self, templates_dir: Optional[str] = None):
        """Initialize Word engine"""
        self.templates_dir = templates_dir or settings.templates_dir
        self.temp_dir = settings.temp_dir
        
        # Template field mappings
        self._field_mappings = {
            # Basic meeting information
            '{{MEETING_TITLE}}': 'meeting_title',
            '{{MEETING_DATE}}': 'meeting_date',
            '{{MEETING_TIME}}': 'meeting_time',
            '{{MEETING_LOCATION}}': 'meeting_location',
            '{{MEETING_DURATION}}': 'meeting_duration',
            '{{MEETING_ORGANIZER}}': 'meeting_organizer',
            '{{GENERATION_DATE}}': 'generation_date',
            
            # Attendees
            '{{ATTENDEES_LIST}}': 'attendees',
            '{{ATTENDEES_COUNT}}': 'attendees_count',
            
            # Agenda and discussions
            '{{AGENDA_ITEMS}}': 'agenda',
            '{{DISCUSSION_TOPICS}}': 'agenda',
            
            # Action items
            '{{ACTION_ITEMS}}': 'action_items',
            '{{PENDING_ACTIONS}}': 'action_items',
            
            # Decisions
            '{{DECISIONS_MADE}}': 'decisions',
            '{{KEY_DECISIONS}}': 'decisions',
            
            # Outcomes
            '{{KEY_OUTCOMES}}': 'key_outcomes',
            '{{MEETING_SUMMARY}}': 'key_outcomes',
            
            # Additional content
            '{{ADDITIONAL_NOTES}}': 'additional_notes',
            '{{NEXT_MEETING}}': 'next_meeting',
        }
        
        if not DOCX_AVAILABLE:
            logger.warning("⚠️ python-docx not available - using text-based processing")
        
        logger.info(f"📝 Word Engine initialized - templates dir: {self.templates_dir}")
    
    async def process_template(
        self,
        template_path: str,
        template_data: Dict[str, Any],
        options: Optional[Dict] = None
    ) -> ProcessingResult:
        """
        Process Word template with provided data
        
        Args:
            template_path: Path to Word template file
            template_data: Data to inject into template
            options: Processing options
            
        Returns:
            ProcessingResult with processed document information
        """
        start_time = datetime.utcnow()
        options = options or {}
        
        if not os.path.exists(template_path):
            return ProcessingResult(
                success=False,
                error_message=f"Template file not found: {template_path}"
            )
        
        logger.info(f"📄 Processing Word template: {os.path.basename(template_path)}")
        
        try:
            if DOCX_AVAILABLE:
                result = await self._process_docx_template(template_path, template_data, options)
            else:
                result = await self._process_text_template(template_path, template_data, options)
            
            # Calculate processing time
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            result.processing_time = processing_time
            
            if result.success:
                logger.info(f"✅ Template processed successfully in {processing_time:.2f}s")
            else:
                logger.error(f"❌ Template processing failed: {result.error_message}")
            
            return result
            
        except Exception as e:
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            logger.error(f"❌ Template processing error: {e}")
            
            return ProcessingResult(
                success=False,
                error_message=str(e),
                processing_time=processing_time
            )
    
    async def _process_docx_template(
        self,
        template_path: str,
        template_data: Dict[str, Any],
        options: Dict
    ) -> ProcessingResult:
        """Process DOCX template using python-docx"""
        try:
            # Load template document
            doc = Document(template_path)
            
            # Track processed fields
            processed_fields = []
            warnings = []
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                if self._contains_placeholder(original_text):
                    new_text = self._replace_placeholders(original_text, template_data)
                    if new_text != original_text:
                        paragraph.text = new_text
                        processed_fields.extend(self._extract_processed_fields(original_text))
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text
                            if self._contains_placeholder(original_text):
                                new_text = self._replace_placeholders(original_text, template_data)
                                if new_text != original_text:
                                    paragraph.text = new_text
                                    processed_fields.extend(self._extract_processed_fields(original_text))
            
            # Handle special data types (tables, lists)
            await self._process_special_content(doc, template_data, options)
            
            # Generate output filename
            output_filename = self._generate_output_filename(template_path, options)
            output_path = os.path.join(self.temp_dir, output_filename)
            
            # Save processed document
            doc.save(output_path)
            
            return ProcessingResult(
                success=True,
                output_path=output_path,
                processed_fields=list(set(processed_fields)),
                warnings=warnings,
                metadata={
                    'template_file': os.path.basename(template_path),
                    'output_file': output_filename,
                    'fields_processed': len(set(processed_fields))
                }
            )
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                error_message=f"DOCX processing error: {str(e)}"
            )
    
    async def _process_text_template(
        self,
        template_path: str,
        template_data: Dict[str, Any],
        options: Dict
    ) -> ProcessingResult:
        """Fallback text-based template processing"""
        try:
            # Read template as text
            with open(template_path, 'r', encoding='utf-8') as f:
                template_content = f.read()
            
            # Replace placeholders
            processed_content = self._replace_placeholders(template_content, template_data)
            processed_fields = self._extract_processed_fields(template_content)
            
            # Generate output filename
            output_filename = self._generate_output_filename(template_path, options, '.txt')
            output_path = os.path.join(self.temp_dir, output_filename)
            
            # Save processed content
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(processed_content)
            
            return ProcessingResult(
                success=True,
                output_path=output_path,
                processed_fields=processed_fields,
                warnings=["Processed as text file - DOCX features not available"],
                metadata={
                    'template_file': os.path.basename(template_path),
                    'output_file': output_filename,
                    'processing_mode': 'text'
                }
            )
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                error_message=f"Text processing error: {str(e)}"
            )
    
    async def _process_special_content(
        self,
        doc: 'Document',
        template_data: Dict[str, Any],
        options: Dict
    ):
        """Process special content like tables and lists"""
        if not DOCX_AVAILABLE:
            return
        
        try:
            # Process attendees table
            if 'attendees' in template_data:
                await self._insert_attendees_table(doc, template_data['attendees'])
            
            # Process agenda items
            if 'agenda' in template_data:
                await self._insert_agenda_items(doc, template_data['agenda'])
            
            # Process action items table
            if 'action_items' in template_data:
                await self._insert_action_items_table(doc, template_data['action_items'])
            
            # Process decisions
            if 'decisions' in template_data:
                await self._insert_decisions_list(doc, template_data['decisions'])
                
        except Exception as e:
            logger.warning(f"⚠️ Special content processing warning: {e}")
    
    async def _insert_attendees_table(self, doc: 'Document', attendees: List[Dict]):
        """Insert attendees table"""
        # Find placeholder paragraph
        placeholder_para = None
        for para in doc.paragraphs:
            if '{{ATTENDEES_TABLE}}' in para.text:
                placeholder_para = para
                break
        
        if not placeholder_para or not attendees:
            return
        
        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = '姓名'
        header_cells[1].text = '職位'
        header_cells[2].text = '單位'
        header_cells[3].text = '出席狀態'
        
        # Add attendees
        for attendee in attendees:
            row_cells = table.add_row().cells
            row_cells[0].text = attendee.get('name', '')
            row_cells[1].text = attendee.get('role', '')
            row_cells[2].text = attendee.get('organization', '')
            row_cells[3].text = attendee.get('present', '出席')
        
        # Insert table before placeholder
        placeholder_para.insert_paragraph_before().element.addnext(table.element)
        
        # Remove placeholder
        placeholder_para.clear()
    
    async def _insert_agenda_items(self, doc: 'Document', agenda: List[Dict]):
        """Insert agenda items"""
        placeholder_para = None
        for para in doc.paragraphs:
            if '{{AGENDA_DETAILS}}' in para.text:
                placeholder_para = para
                break
        
        if not placeholder_para or not agenda:
            return
        
        # Insert agenda items
        for i, item in enumerate(agenda, 1):
            # Add agenda item title
            title_para = placeholder_para.insert_paragraph_before(f"{i}. {item.get('title', '')}")
            title_para.style = 'Heading 3'
            
            # Add description if available
            if item.get('description'):
                desc_para = placeholder_para.insert_paragraph_before(f"   說明: {item['description']}")
            
            # Add key points
            if item.get('key_points'):
                points_para = placeholder_para.insert_paragraph_before("   討論要點:")
                for point in item['key_points']:
                    placeholder_para.insert_paragraph_before(f"     • {point}")
        
        # Remove placeholder
        placeholder_para.clear()
    
    async def _insert_action_items_table(self, doc: 'Document', action_items: List[Dict]):
        """Insert action items table"""
        placeholder_para = None
        for para in doc.paragraphs:
            if '{{ACTION_ITEMS_TABLE}}' in para.text:
                placeholder_para = para
                break
        
        if not placeholder_para or not action_items:
            return
        
        # Create action items table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header
        headers = ['項目', '負責人', '截止日期', '優先級', '狀態']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        # Add action items
        for action in action_items:
            row_cells = table.add_row().cells
            row_cells[0].text = action.get('task', '')
            row_cells[1].text = action.get('assignee', '')
            row_cells[2].text = action.get('due_date', '')
            row_cells[3].text = action.get('priority', '')
            row_cells[4].text = action.get('status', '待處理')
        
        # Insert table
        placeholder_para.insert_paragraph_before().element.addnext(table.element)
        placeholder_para.clear()
    
    async def _insert_decisions_list(self, doc: 'Document', decisions: List[Dict]):
        """Insert decisions list"""
        placeholder_para = None
        for para in doc.paragraphs:
            if '{{DECISIONS_LIST}}' in para.text:
                placeholder_para = para
                break
        
        if not placeholder_para or not decisions:
            return
        
        # Insert decisions
        for i, decision in enumerate(decisions, 1):
            decision_para = placeholder_para.insert_paragraph_before(f"{i}. {decision.get('decision', '')}")
            decision_para.style = 'List Number'
            
            if decision.get('rationale'):
                placeholder_para.insert_paragraph_before(f"   理由: {decision['rationale']}")
            
            if decision.get('responsible_party'):
                placeholder_para.insert_paragraph_before(f"   負責單位: {decision['responsible_party']}")
        
        placeholder_para.clear()
    
    def _contains_placeholder(self, text: str) -> bool:
        """Check if text contains template placeholders"""
        return any(placeholder in text for placeholder in self._field_mappings.keys())
    
    def _replace_placeholders(self, text: str, template_data: Dict[str, Any]) -> str:
        """Replace placeholders in text with actual data"""
        result = text
        
        for placeholder, data_key in self._field_mappings.items():
            if placeholder in result:
                # Handle nested data
                value = self._get_nested_value(template_data, data_key)
                
                if value is not None:
                    # Format value based on type
                    if isinstance(value, list):
                        if data_key == 'key_outcomes':
                            formatted_value = '\n'.join(f"• {item}" for item in value)
                        elif data_key == 'attendees':
                            formatted_value = ', '.join(
                                attendee.get('name', '') for attendee in value 
                                if isinstance(attendee, dict)
                            )
                            if not formatted_value:
                                formatted_value = '\n'.join(str(item) for item in value)
                        else:
                            formatted_value = ', '.join(str(item) for item in value)
                    elif isinstance(value, dict):
                        # Handle next meeting info
                        if data_key == 'next_meeting':
                            parts = []
                            if value.get('date'):
                                parts.append(f"日期: {value['date']}")
                            if value.get('time'):
                                parts.append(f"時間: {value['time']}")
                            if value.get('location'):
                                parts.append(f"地點: {value['location']}")
                            formatted_value = ', '.join(parts) if parts else ''
                        else:
                            formatted_value = str(value)
                    else:
                        formatted_value = str(value)
                    
                    result = result.replace(placeholder, formatted_value)
                else:
                    # Replace with empty string if no data
                    result = result.replace(placeholder, '')
        
        return result
    
    def _get_nested_value(self, data: Dict, key: str) -> Any:
        """Get nested value from data dictionary"""
        keys = key.split('.')
        value = data
        
        for k in keys:
            if isinstance(value, dict) and k in value:
                value = value[k]
            else:
                return None
        
        return value
    
    def _extract_processed_fields(self, text: str) -> List[str]:
        """Extract field names that were processed"""
        fields = []
        for placeholder, field_name in self._field_mappings.items():
            if placeholder in text:
                fields.append(field_name)
        return fields
    
    def _generate_output_filename(
        self, 
        template_path: str, 
        options: Dict,
        extension: str = '.docx'
    ) -> str:
        """Generate output filename"""
        template_name = Path(template_path).stem
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        
        custom_name = options.get('output_filename')
        if custom_name:
            name = Path(custom_name).stem
        else:
            name = f"{template_name}_processed_{timestamp}"
        
        return f"{name}{extension}"
    
    async def validate_template(self, template_path: str) -> Dict[str, Any]:
        """Validate template file and extract field information"""
        if not os.path.exists(template_path):
            return {
                'valid': False,
                'error': 'Template file not found'
            }
        
        try:
            if DOCX_AVAILABLE and template_path.endswith('.docx'):
                doc = Document(template_path)
                
                # Extract all text
                all_text = []
                for para in doc.paragraphs:
                    all_text.append(para.text)
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                all_text.append(para.text)
                
                full_text = '\n'.join(all_text)
                
            else:
                # Text-based validation
                with open(template_path, 'r', encoding='utf-8') as f:
                    full_text = f.read()
            
            # Find placeholders
            found_placeholders = []
            for placeholder in self._field_mappings.keys():
                if placeholder in full_text:
                    found_placeholders.append(placeholder)
            
            return {
                'valid': True,
                'placeholders_found': found_placeholders,
                'total_placeholders': len(found_placeholders),
                'template_size': os.path.getsize(template_path),
                'supported_fields': list(self._field_mappings.values())
            }
            
        except Exception as e:
            return {
                'valid': False,
                'error': str(e)
            }
    
    def get_supported_fields(self) -> Dict[str, str]:
        """Get all supported template fields"""
        return self._field_mappings.copy()
    
    async def create_sample_template(self, template_type: str = 'standard') -> str:
        """Create a sample template file"""
        if template_type == 'standard':
            content = self._get_standard_template_content()
        elif template_type == 'executive':
            content = self._get_executive_template_content()
        else:
            content = self._get_standard_template_content()
        
        # Save sample template
        filename = f"sample_{template_type}_template.docx"
        output_path = os.path.join(self.templates_dir, filename)
        
        if DOCX_AVAILABLE:
            doc = Document()
            
            # Add content paragraphs
            for line in content.split('\n'):
                if line.strip():
                    para = doc.add_paragraph(line)
                    if line.startswith('{{') and line.endswith('}}'):
                        # Style placeholders differently
                        para.runs[0].bold = True
            
            doc.save(output_path)
        else:
            # Save as text file
            output_path = output_path.replace('.docx', '.txt')
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
        
        logger.info(f"📝 Sample template created: {output_path}")
        return output_path
    
    def _get_standard_template_content(self) -> str:
        """Get standard template content"""
        return """會議記錄

會議標題: {{MEETING_TITLE}}
會議日期: {{MEETING_DATE}}
會議時間: {{MEETING_TIME}}
會議地點: {{MEETING_LOCATION}}
會議時長: {{MEETING_DURATION}}
主辦人: {{MEETING_ORGANIZER}}

出席人員:
{{ATTENDEES_TABLE}}

議程討論:
{{AGENDA_DETAILS}}

行動項目:
{{ACTION_ITEMS_TABLE}}

決議事項:
{{DECISIONS_LIST}}

重要結論:
{{KEY_OUTCOMES}}

其他事項:
{{ADDITIONAL_NOTES}}

下次會議:
{{NEXT_MEETING}}

記錄產生時間: {{GENERATION_DATE}}"""
    
    def _get_executive_template_content(self) -> str:
        """Get executive template content"""
        return """高階主管會議記錄

會議標題: {{MEETING_TITLE}}
會議日期: {{MEETING_DATE}}
與會主管: {{ATTENDEES_LIST}}

重要決議:
{{DECISIONS_LIST}}

策略方向:
{{KEY_OUTCOMES}}

後續行動:
{{ACTION_ITEMS_TABLE}}

記錄時間: {{GENERATION_DATE}}"""