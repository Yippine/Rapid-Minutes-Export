import logging
import tempfile
import os
from typing import Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path

# Optional magic import for MIME type detection
try:
    import magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False

logger = logging.getLogger(__name__)


class DocumentValidator:
    """Comprehensive document validation service following SESE principles"""
    
    def __init__(self):
        self.supported_formats = {
            "docx": {
                "mime_types": ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
                "magic_bytes": [b"PK"],  # ZIP-based format
                "extensions": [".docx"]
            },
            "pdf": {
                "mime_types": ["application/pdf"],
                "magic_bytes": [b"%PDF-"],
                "extensions": [".pdf"]
            },
            "txt": {
                "mime_types": ["text/plain"],
                "magic_bytes": [],  # Variable
                "extensions": [".txt"]
            }
        }
    
    def validate_document_comprehensive(self, document_content: bytes, 
                                      expected_format: str, 
                                      source_filename: Optional[str] = None) -> Dict[str, Any]:
        """Comprehensive document validation implementing ICE principle - exhaustive coverage"""
        
        validation_result = {
            "is_valid": False,
            "format": expected_format.lower(),
            "file_size": len(document_content),
            "source_filename": source_filename,
            "validation_timestamp": datetime.now().isoformat(),
            "checks_performed": [],
            "issues": [],
            "warnings": [],
            "quality_metrics": {},
            "recommendations": []
        }
        
        try:
            # 1. Basic validation checks
            self._perform_basic_validation(document_content, validation_result)
            
            # 2. Format-specific validation
            self._perform_format_validation(document_content, expected_format, validation_result)
            
            # 3. Content quality validation
            self._perform_content_validation(document_content, expected_format, validation_result)
            
            # 4. Security validation
            self._perform_security_validation(document_content, expected_format, validation_result)
            
            # 5. Generate quality score and recommendations
            self._calculate_quality_metrics(validation_result)
            
            # 6. Determine overall validity
            validation_result["is_valid"] = (
                len(validation_result["issues"]) == 0 and
                validation_result["quality_metrics"].get("overall_score", 0) >= 70
            )
            
            logger.info(f"Document validation completed. Valid: {validation_result['is_valid']}")
            return validation_result
            
        except Exception as e:
            logger.error(f"Document validation failed: {e}")
            validation_result["issues"].append(f"Validation error: {str(e)}")
            return validation_result
    
    def _perform_basic_validation(self, document_content: bytes, result: Dict[str, Any]):
        """Perform basic document validation checks"""
        result["checks_performed"].append("basic_validation")
        
        # Size validation
        if len(document_content) == 0:
            result["issues"].append("Document is empty")
            return
        
        # Minimum size check
        if len(document_content) < 100:
            result["warnings"].append("Document is very small (< 100 bytes)")
        
        # Maximum size check (50MB limit)
        max_size = 50 * 1024 * 1024
        if len(document_content) > max_size:
            result["issues"].append(f"Document exceeds maximum size limit ({max_size} bytes)")
        
        # Basic corruption check
        try:
            # Check if content is readable
            if len(document_content) > 1000:
                # Sample different parts of the document
                samples = [
                    document_content[:100],
                    document_content[len(document_content)//2:len(document_content)//2+100],
                    document_content[-100:]
                ]
                for i, sample in enumerate(samples):
                    if sample == b'\x00' * len(sample):
                        result["warnings"].append(f"Document section {i+1} appears to contain null bytes")
        except Exception as e:
            result["warnings"].append(f"Basic corruption check failed: {str(e)}")
    
    def _perform_format_validation(self, document_content: bytes, expected_format: str, result: Dict[str, Any]):
        """Perform format-specific validation"""
        result["checks_performed"].append("format_validation")
        
        expected_format = expected_format.lower()
        
        if expected_format not in self.supported_formats:
            result["issues"].append(f"Unsupported format: {expected_format}")
            return
        
        format_spec = self.supported_formats[expected_format]
        
        # Magic bytes validation
        if format_spec["magic_bytes"]:
            format_validated = False
            for magic_byte in format_spec["magic_bytes"]:
                if document_content.startswith(magic_byte):
                    format_validated = True
                    break
            
            if not format_validated:
                result["issues"].append(f"Document does not start with expected magic bytes for {expected_format}")
        
        # Format-specific validation
        if expected_format == "docx":
            self._validate_docx_format(document_content, result)
        elif expected_format == "pdf":
            self._validate_pdf_format(document_content, result)
        elif expected_format == "txt":
            self._validate_txt_format(document_content, result)
    
    def _validate_docx_format(self, document_content: bytes, result: Dict[str, Any]):
        """Validate DOCX document format"""
        try:
            # DOCX is a ZIP file, check ZIP structure
            if not document_content.startswith(b'PK'):
                result["issues"].append("DOCX file does not have valid ZIP header")
                return
            
            # Try to extract and validate using temporary file
            with tempfile.TemporaryDirectory() as temp_dir:
                docx_path = os.path.join(temp_dir, "test.docx")
                with open(docx_path, "wb") as f:
                    f.write(document_content)
                
                try:
                    from docx import Document
                    doc = Document(docx_path)
                    
                    # Validate document structure
                    paragraph_count = len(doc.paragraphs)
                    table_count = len(doc.tables)
                    
                    result["quality_metrics"]["paragraph_count"] = paragraph_count
                    result["quality_metrics"]["table_count"] = table_count
                    
                    if paragraph_count == 0:
                        result["warnings"].append("Document has no paragraphs")
                    
                    # Check for essential content
                    text_content = "\n".join([p.text for p in doc.paragraphs])
                    if len(text_content.strip()) == 0:
                        result["issues"].append("Document appears to have no readable text content")
                    
                except Exception as e:
                    result["issues"].append(f"Failed to parse DOCX content: {str(e)}")
                    
        except Exception as e:
            result["issues"].append(f"DOCX format validation failed: {str(e)}")
    
    def _validate_pdf_format(self, document_content: bytes, result: Dict[str, Any]):
        """Validate PDF document format"""
        try:
            # Check PDF header
            if not document_content.startswith(b'%PDF-'):
                result["issues"].append("Invalid PDF header")
                return
            
            # Check PDF version
            header_line = document_content.split(b'\n')[0]
            try:
                version = header_line.decode('ascii').replace('%PDF-', '')
                result["quality_metrics"]["pdf_version"] = version
                
                if float(version) < 1.4:
                    result["warnings"].append(f"PDF version {version} is quite old")
            except:
                result["warnings"].append("Could not determine PDF version")
            
            # Check for PDF trailer
            if b'%%EOF' not in document_content[-200:]:
                result["issues"].append("PDF does not have proper trailer")
            
            # Check for basic PDF structure
            required_elements = [b'/Type', b'/Pages', b'obj']
            missing_elements = []
            for element in required_elements:
                if element not in document_content:
                    missing_elements.append(element.decode())
            
            if missing_elements:
                result["warnings"].append(f"PDF may be missing elements: {missing_elements}")
                
        except Exception as e:
            result["issues"].append(f"PDF format validation failed: {str(e)}")
    
    def _validate_txt_format(self, document_content: bytes, result: Dict[str, Any]):
        """Validate text document format"""
        try:
            # Try to decode as text
            try:
                text_content = document_content.decode('utf-8')
                result["quality_metrics"]["encoding"] = "utf-8"
            except UnicodeDecodeError:
                try:
                    text_content = document_content.decode('latin-1')
                    result["quality_metrics"]["encoding"] = "latin-1"
                    result["warnings"].append("Document uses latin-1 encoding instead of UTF-8")
                except:
                    result["issues"].append("Document cannot be decoded as text")
                    return
            
            # Basic text quality checks
            line_count = len(text_content.splitlines())
            word_count = len(text_content.split())
            char_count = len(text_content)
            
            result["quality_metrics"]["line_count"] = line_count
            result["quality_metrics"]["word_count"] = word_count
            result["quality_metrics"]["char_count"] = char_count
            
            if word_count == 0:
                result["issues"].append("Text document appears to have no words")
            elif word_count < 10:
                result["warnings"].append("Text document has very few words")
                
        except Exception as e:
            result["issues"].append(f"Text format validation failed: {str(e)}")
    
    def _perform_content_validation(self, document_content: bytes, expected_format: str, result: Dict[str, Any]):
        """Perform content quality validation"""
        result["checks_performed"].append("content_validation")
        
        try:
            # Format-specific content validation
            if expected_format == "docx":
                self._validate_docx_content(document_content, result)
            elif expected_format == "pdf":
                self._validate_pdf_content(document_content, result)
            elif expected_format == "txt":
                self._validate_txt_content(document_content, result)
                
        except Exception as e:
            result["warnings"].append(f"Content validation failed: {str(e)}")
    
    def _validate_docx_content(self, document_content: bytes, result: Dict[str, Any]):
        """Validate DOCX content quality"""
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                docx_path = os.path.join(temp_dir, "test.docx")
                with open(docx_path, "wb") as f:
                    f.write(document_content)
                
                from docx import Document
                doc = Document(docx_path)
                
                # Check for meeting minutes structure
                headings = []
                for paragraph in doc.paragraphs:
                    if paragraph.style.name.startswith('Heading'):
                        headings.append(paragraph.text)
                
                result["quality_metrics"]["heading_count"] = len(headings)
                
                # Look for expected meeting minutes sections
                expected_sections = ["meeting", "attendee", "topic", "decision", "action"]
                found_sections = []
                
                all_text = " ".join([p.text.lower() for p in doc.paragraphs])
                for section in expected_sections:
                    if section in all_text:
                        found_sections.append(section)
                
                result["quality_metrics"]["found_meeting_sections"] = found_sections
                
                if len(found_sections) < 3:
                    result["warnings"].append("Document may not contain typical meeting minutes content")
                    
        except Exception as e:
            result["warnings"].append(f"DOCX content validation failed: {str(e)}")
    
    def _validate_pdf_content(self, document_content: bytes, result: Dict[str, Any]):
        """Validate PDF content quality"""
        # Basic PDF content validation
        try:
            # Check for text content indicators
            text_indicators = [b'/Text', b'/Font', b'BT', b'ET']  # Basic text operators
            text_found = sum(1 for indicator in text_indicators if indicator in document_content)
            
            result["quality_metrics"]["text_indicators_found"] = text_found
            
            if text_found == 0:
                result["warnings"].append("PDF may not contain readable text content")
                
        except Exception as e:
            result["warnings"].append(f"PDF content validation failed: {str(e)}")
    
    def _validate_txt_content(self, document_content: bytes, result: Dict[str, Any]):
        """Validate text content quality"""
        try:
            text_content = document_content.decode('utf-8', errors='ignore')
            
            # Check for meeting-related keywords
            meeting_keywords = ["meeting", "agenda", "minutes", "discussion", "decision", "action", "attendees"]
            found_keywords = [kw for kw in meeting_keywords if kw.lower() in text_content.lower()]
            
            result["quality_metrics"]["meeting_keywords_found"] = found_keywords
            
            if len(found_keywords) < 2:
                result["warnings"].append("Text may not be meeting-related content")
                
        except Exception as e:
            result["warnings"].append(f"Text content validation failed: {str(e)}")
    
    def _perform_security_validation(self, document_content: bytes, expected_format: str, result: Dict[str, Any]):
        """Perform basic security validation"""
        result["checks_performed"].append("security_validation")
        
        try:
            # Check for suspicious patterns
            suspicious_patterns = [
                b'<script',
                b'javascript:',
                b'vbscript:',
                b'<?php',
                b'<%',
                b'eval(',
                b'exec('
            ]
            
            found_suspicious = []
            for pattern in suspicious_patterns:
                if pattern in document_content.lower():
                    found_suspicious.append(pattern.decode('utf-8', errors='ignore'))
            
            if found_suspicious:
                result["warnings"].append(f"Document contains potentially suspicious patterns: {found_suspicious}")
            
            # File size based security check
            if len(document_content) > 10 * 1024 * 1024:  # 10MB
                result["warnings"].append("Document is unusually large")
                
        except Exception as e:
            result["warnings"].append(f"Security validation failed: {str(e)}")
    
    def _calculate_quality_metrics(self, result: Dict[str, Any]):
        """Calculate overall quality metrics and generate recommendations"""
        try:
            # Base score
            quality_score = 100
            
            # Deduct points for issues and warnings
            quality_score -= len(result["issues"]) * 25
            quality_score -= len(result["warnings"]) * 5
            
            # Bonus points for comprehensive content
            content_bonus = 0
            if "meeting_keywords_found" in result["quality_metrics"]:
                content_bonus += len(result["quality_metrics"]["meeting_keywords_found"]) * 2
            
            if "found_meeting_sections" in result["quality_metrics"]:
                content_bonus += len(result["quality_metrics"]["found_meeting_sections"]) * 3
            
            quality_score += min(content_bonus, 20)  # Cap bonus at 20 points
            
            # Ensure score is within bounds
            result["quality_metrics"]["overall_score"] = max(0, min(100, quality_score))
            
            # Generate recommendations
            if len(result["issues"]) > 0:
                result["recommendations"].append("Address critical format issues before using the document")
            
            if len(result["warnings"]) > 3:
                result["recommendations"].append("Review and address document quality warnings")
            
            if result["quality_metrics"]["overall_score"] < 70:
                result["recommendations"].append("Document quality is below recommended standards")
            
            if result["quality_metrics"]["overall_score"] >= 90:
                result["recommendations"].append("Document meets high quality standards")
                
        except Exception as e:
            logger.error(f"Quality metrics calculation failed: {e}")
            result["quality_metrics"]["overall_score"] = 50  # Default fallback score
    
    def get_format_requirements(self, format_type: str) -> Dict[str, Any]:
        """Get format requirements and validation criteria"""
        format_type = format_type.lower()
        
        if format_type in self.supported_formats:
            format_spec = self.supported_formats[format_type].copy()
            format_spec["validation_criteria"] = self._get_format_criteria(format_type)
            return format_spec
        else:
            return {"error": f"Unsupported format: {format_type}"}
    
    def _get_format_criteria(self, format_type: str) -> Dict[str, Any]:
        """Get specific validation criteria for each format"""
        criteria = {
            "docx": {
                "min_size": 1024,  # 1KB
                "max_size": 50 * 1024 * 1024,  # 50MB
                "required_structure": ["paragraphs", "styles"],
                "quality_indicators": ["headings", "tables", "text_content"]
            },
            "pdf": {
                "min_size": 500,  # 500 bytes
                "max_size": 50 * 1024 * 1024,  # 50MB
                "required_structure": ["header", "trailer", "objects"],
                "quality_indicators": ["text_content", "fonts", "pages"]
            },
            "txt": {
                "min_size": 10,  # 10 bytes
                "max_size": 10 * 1024 * 1024,  # 10MB
                "required_structure": ["readable_text"],
                "quality_indicators": ["word_count", "line_count", "encoding"]
            }
        }
        
        return criteria.get(format_type, {})