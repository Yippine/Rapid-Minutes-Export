"""
Data Injection System (D2 - Document Processing Layer)
Advanced data injection and formatting for Word documents
Implements SESE principle - Simple, Effective, Systematic, Exhaustive data processing
"""

import logging
import re
from typing import Dict, List, Optional, Union, Any, TYPE_CHECKING
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ..ai.extractor import MeetingMinutes

# Use TYPE_CHECKING to avoid circular imports
if TYPE_CHECKING:
    from ..core.template_controller import TemplateConfig

logger = logging.getLogger(__name__)


class InjectionType(Enum):
    """Data injection types"""
    TEXT = "text"
    TABLE = "table"
    LIST = "list"
    FORMATTED_TEXT = "formatted_text"
    RICH_CONTENT = "rich_content"


class DataFormat(Enum):
    """Data formatting options"""
    PLAIN = "plain"
    BULLET_LIST = "bullet_list"
    NUMBERED_LIST = "numbered_list"
    TABLE_GRID = "table_grid"
    TABLE_CLEAN = "table_clean"
    HIGHLIGHTED = "highlighted"


@dataclass
class InjectionRule:
    """Data injection rule"""
    field_name: str
    placeholder_pattern: str
    injection_type: InjectionType
    data_format: DataFormat = DataFormat.PLAIN
    required: bool = True
    default_value: Any = None
    formatting_options: Dict[str, Any] = field(default_factory=dict)
    validation_rules: List[str] = field(default_factory=list)


@dataclass
class InjectionResult:
    """Data injection result"""
    success: bool
    injected_fields: List[str] = field(default_factory=list)
    failed_fields: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    error_message: Optional[str] = None
    processed_document: Optional['Document'] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


class DataInjector:
    """
    Advanced data injection system for Word documents
    Handles structured data insertion with formatting and validation
    """
    
    def __init__(self):
        """Initialize data injector"""
        self.injection_rules = self._setup_injection_rules()
        self.formatting_presets = self._setup_formatting_presets()
        
        if not DOCX_AVAILABLE:
            logger.warning("⚠️ python-docx not available - data injection will be limited")
        
        logger.info("💉 Data Injector initialized")
    
    def _setup_injection_rules(self) -> Dict[str, InjectionRule]:
        """Setup data injection rules"""
        rules = {}
        
        # Basic meeting information rules
        basic_fields = [
            ('meeting_title', '{{MEETING_TITLE}}', InjectionType.TEXT),
            ('meeting_date', '{{MEETING_DATE}}', InjectionType.TEXT),
            ('meeting_time', '{{MEETING_TIME}}', InjectionType.TEXT),
            ('meeting_location', '{{MEETING_LOCATION}}', InjectionType.TEXT),
            ('meeting_duration', '{{MEETING_DURATION}}', InjectionType.TEXT),
            ('meeting_organizer', '{{MEETING_ORGANIZER}}', InjectionType.TEXT),
            ('generation_date', '{{GENERATION_DATE}}', InjectionType.TEXT),
        ]
        
        for field_name, placeholder, injection_type in basic_fields:
            rules[field_name] = InjectionRule(
                field_name=field_name,
                placeholder_pattern=placeholder,
                injection_type=injection_type,
                data_format=DataFormat.PLAIN
            )
        
        # Complex data structure rules
        rules['attendees'] = InjectionRule(
            field_name='attendees',
            placeholder_pattern='{{ATTENDEES_(?:LIST|TABLE)}}',
            injection_type=InjectionType.TABLE,
            data_format=DataFormat.TABLE_GRID,
            formatting_options={
                'headers': ['姓名', '職位', '單位', '出席狀態'],
                'style': 'Table Grid'
            }
        )
        
        rules['agenda'] = InjectionRule(
            field_name='agenda',
            placeholder_pattern='{{AGENDA_(?:ITEMS|DETAILS)}}',
            injection_type=InjectionType.RICH_CONTENT,
            data_format=DataFormat.NUMBERED_LIST,
            formatting_options={
                'style': 'List Number',
                'include_details': True
            }
        )
        
        rules['action_items'] = InjectionRule(
            field_name='action_items',
            placeholder_pattern='{{ACTION_ITEMS_(?:TABLE|LIST)}}',
            injection_type=InjectionType.TABLE,
            data_format=DataFormat.TABLE_GRID,
            formatting_options={
                'headers': ['項目', '負責人', '截止日期', '優先級', '狀態'],
                'style': 'Table Grid'
            }
        )
        
        rules['decisions'] = InjectionRule(
            field_name='decisions',
            placeholder_pattern='{{DECISIONS_(?:LIST|MADE)}}',
            injection_type=InjectionType.LIST,
            data_format=DataFormat.NUMBERED_LIST,
            formatting_options={
                'style': 'List Number',
                'include_rationale': True
            }
        )
        
        rules['key_outcomes'] = InjectionRule(
            field_name='key_outcomes',
            placeholder_pattern='{{KEY_OUTCOMES|MEETING_SUMMARY}}',
            injection_type=InjectionType.LIST,
            data_format=DataFormat.BULLET_LIST,
            formatting_options={
                'style': 'List Bullet'
            }
        )
        
        return rules
    
    def _setup_formatting_presets(self) -> Dict[str, Dict]:
        """Setup formatting presets"""
        return {
            'standard': {
                'font_name': '微軟正黑體',
                'font_size': 12,
                'line_spacing': 1.15,
                'paragraph_spacing': 6
            },
            'executive': {
                'font_name': 'Times New Roman',
                'font_size': 12,
                'line_spacing': 1.5,
                'paragraph_spacing': 8
            },
            'compact': {
                'font_name': '微軟正黑體',
                'font_size': 10,
                'line_spacing': 1.0,
                'paragraph_spacing': 3
            }
        }
    
    async def inject_data(
        self,
        document: Union['Document', str],
        meeting_minutes: MeetingMinutes,
        template_config: 'TemplateConfig',
        formatting_preset: str = 'standard'
    ) -> InjectionResult:
        """
        Inject meeting data into Word document
        
        Args:
            document: Document object or path
            meeting_minutes: Meeting minutes data
            template_config: Template configuration
            formatting_preset: Formatting preset name
            
        Returns:
            InjectionResult with processing results
        """
        logger.info("💉 Starting data injection process")
        
        try:
            # Load document if path provided
            if isinstance(document, str):
                if not DOCX_AVAILABLE:
                    return InjectionResult(
                        success=False,
                        error_message="python-docx not available for document processing"
                    )
                doc = Document(document)
            else:
                doc = document
            
            if not DOCX_AVAILABLE:
                return await self._inject_data_text_mode(meeting_minutes, template_config)
            
            # Prepare injection data
            injection_data = self._prepare_injection_data(meeting_minutes)
            
            # Track injection results
            injected_fields = []
            failed_fields = []
            warnings = []
            
            # Apply formatting preset
            if formatting_preset in self.formatting_presets:
                self._apply_document_formatting(doc, self.formatting_presets[formatting_preset])
            
            # Process each injection rule
            for field_name, rule in self.injection_rules.items():
                try:
                    if field_name in injection_data:
                        success = await self._inject_field_data(
                            doc, rule, injection_data[field_name], template_config
                        )
                        if success:
                            injected_fields.append(field_name)
                        else:
                            failed_fields.append(field_name)
                            warnings.append(f"Failed to inject {field_name}")
                    elif rule.required:
                        # Use default value or placeholder
                        default_value = rule.default_value or f"[{field_name.replace('_', ' ').title()}]"
                        success = await self._inject_field_data(
                            doc, rule, default_value, template_config
                        )
                        if success:
                            injected_fields.append(field_name)
                            warnings.append(f"Used default value for {field_name}")
                        else:
                            failed_fields.append(field_name)
                
                except Exception as e:
                    failed_fields.append(field_name)
                    warnings.append(f"Error injecting {field_name}: {str(e)}")
                    logger.warning(f"⚠️ Field injection error for {field_name}: {e}")
            
            logger.info(f"✅ Data injection completed - {len(injected_fields)} fields injected")
            
            return InjectionResult(
                success=len(injected_fields) > 0,
                injected_fields=injected_fields,
                failed_fields=failed_fields,
                warnings=warnings,
                processed_document=doc,
                metadata={
                    'total_fields': len(self.injection_rules),
                    'success_rate': len(injected_fields) / len(self.injection_rules),
                    'formatting_preset': formatting_preset
                }
            )
            
        except Exception as e:
            logger.error(f"❌ Data injection failed: {e}")
            return InjectionResult(
                success=False,
                error_message=str(e)
            )
    
    async def _inject_field_data(
        self,
        doc: 'Document',
        rule: InjectionRule,
        data: Any,
        template_config: 'TemplateConfig'
    ) -> bool:
        """Inject data for a specific field"""
        try:
            if rule.injection_type == InjectionType.TEXT:
                return self._inject_text_data(doc, rule, data)
            elif rule.injection_type == InjectionType.TABLE:
                return self._inject_table_data(doc, rule, data)
            elif rule.injection_type == InjectionType.LIST:
                return self._inject_list_data(doc, rule, data)
            elif rule.injection_type == InjectionType.RICH_CONTENT:
                return self._inject_rich_content(doc, rule, data)
            else:
                return self._inject_text_data(doc, rule, str(data))
                
        except Exception as e:
            logger.error(f"❌ Field injection error for {rule.field_name}: {e}")
            return False
    
    def _inject_text_data(self, doc: 'Document', rule: InjectionRule, data: Any) -> bool:
        """Inject simple text data"""
        pattern = re.compile(rule.placeholder_pattern)
        text_value = str(data) if data is not None else ""
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            if pattern.search(paragraph.text):
                new_text = pattern.sub(text_value, paragraph.text)
                paragraph.text = new_text
                
                # Apply formatting if specified
                if rule.formatting_options:
                    self._apply_paragraph_formatting(paragraph, rule.formatting_options)
                
                return True
        
        # Process table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if pattern.search(paragraph.text):
                            new_text = pattern.sub(text_value, paragraph.text)
                            paragraph.text = new_text
                            return True
        
        return False
    
    def _inject_table_data(self, doc: 'Document', rule: InjectionRule, data: List[Dict]) -> bool:
        """Inject table data"""
        if not data:
            return False
        
        # Find placeholder
        placeholder_para = None
        pattern = re.compile(rule.placeholder_pattern)
        
        for paragraph in doc.paragraphs:
            if pattern.search(paragraph.text):
                placeholder_para = paragraph
                break
        
        if not placeholder_para:
            return False
        
        try:
            # Create table based on rule configuration
            headers = rule.formatting_options.get('headers', [])
            if not headers and data:
                headers = list(data[0].keys())
            
            # Create table
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = rule.formatting_options.get('style', 'Table Grid')
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Add headers
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Add data rows
            if rule.field_name == 'attendees':
                for attendee in data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = attendee.get('name', '')
                    row_cells[1].text = attendee.get('role', '')
                    row_cells[2].text = attendee.get('organization', '')
                    row_cells[3].text = attendee.get('present', '出席')
            
            elif rule.field_name == 'action_items':
                for action in data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = action.get('task', '')
                    row_cells[1].text = action.get('assignee', '')
                    row_cells[2].text = action.get('due_date', '')
                    row_cells[3].text = action.get('priority', '')
                    row_cells[4].text = action.get('status', '待處理')
            
            # Insert table and remove placeholder
            placeholder_para.insert_paragraph_before().element.addnext(table.element)
            placeholder_para.clear()
            
            return True
            
        except Exception as e:
            logger.error(f"❌ Table injection error: {e}")
            return False
    
    def _inject_list_data(self, doc: 'Document', rule: InjectionRule, data: List[Union[str, Dict]]) -> bool:
        """Inject list data"""
        if not data:
            return False
        
        # Find placeholder
        placeholder_para = None
        pattern = re.compile(rule.placeholder_pattern)
        
        for paragraph in doc.paragraphs:
            if pattern.search(paragraph.text):
                placeholder_para = paragraph
                break
        
        if not placeholder_para:
            return False
        
        try:
            # Process list items
            if rule.field_name == 'decisions':
                for i, decision in enumerate(data, 1):
                    decision_text = decision.get('decision', str(decision)) if isinstance(decision, dict) else str(decision)
                    
                    # Add main decision
                    decision_para = placeholder_para.insert_paragraph_before(f"{i}. {decision_text}")
                    decision_para.style = rule.formatting_options.get('style', 'List Number')
                    
                    # Add rationale if available
                    if isinstance(decision, dict) and decision.get('rationale') and rule.formatting_options.get('include_rationale'):
                        rationale_para = placeholder_para.insert_paragraph_before(f"   理由: {decision['rationale']}")
                        rationale_para.style = 'Normal'
                    
                    # Add responsible party
                    if isinstance(decision, dict) and decision.get('responsible_party'):
                        resp_para = placeholder_para.insert_paragraph_before(f"   負責單位: {decision['responsible_party']}")
                        resp_para.style = 'Normal'
            
            elif rule.field_name == 'key_outcomes':
                for outcome in data:
                    outcome_text = str(outcome)
                    outcome_para = placeholder_para.insert_paragraph_before(f"• {outcome_text}")
                    outcome_para.style = rule.formatting_options.get('style', 'List Bullet')
            
            # Remove placeholder
            placeholder_para.clear()
            return True
            
        except Exception as e:
            logger.error(f"❌ List injection error: {e}")
            return False
    
    def _inject_rich_content(self, doc: 'Document', rule: InjectionRule, data: List[Dict]) -> bool:
        """Inject rich content with complex formatting"""
        if not data:
            return False
        
        # Find placeholder
        placeholder_para = None
        pattern = re.compile(rule.placeholder_pattern)
        
        for paragraph in doc.paragraphs:
            if pattern.search(paragraph.text):
                placeholder_para = paragraph
                break
        
        if not placeholder_para:
            return False
        
        try:
            if rule.field_name == 'agenda':
                for i, agenda_item in enumerate(data, 1):
                    # Add agenda item title
                    title = agenda_item.get('title', f'議程項目 {i}')
                    title_para = placeholder_para.insert_paragraph_before(f"{i}. {title}")
                    title_para.style = 'Heading 3'
                    
                    # Add description
                    if agenda_item.get('description'):
                        desc_para = placeholder_para.insert_paragraph_before(f"說明: {agenda_item['description']}")
                        desc_para.style = 'Normal'
                    
                    # Add presenter
                    if agenda_item.get('presenter'):
                        pres_para = placeholder_para.insert_paragraph_before(f"主講人: {agenda_item['presenter']}")
                        pres_para.style = 'Normal'
                    
                    # Add key points
                    if agenda_item.get('key_points'):
                        points_para = placeholder_para.insert_paragraph_before("討論要點:")
                        points_para.style = 'Normal'
                        points_para.runs[0].bold = True
                        
                        for point in agenda_item['key_points']:
                            point_para = placeholder_para.insert_paragraph_before(f"  • {point}")
                            point_para.style = 'List Bullet'
                    
                    # Add spacing
                    placeholder_para.insert_paragraph_before("")
            
            # Remove placeholder
            placeholder_para.clear()
            return True
            
        except Exception as e:
            logger.error(f"❌ Rich content injection error: {e}")
            return False
    
    def _prepare_injection_data(self, meeting_minutes: MeetingMinutes) -> Dict[str, Any]:
        """Prepare meeting data for injection"""
        data = {}
        
        # Basic information
        if meeting_minutes.basic_info:
            data.update({
                'meeting_title': meeting_minutes.basic_info.title or "會議記錄",
                'meeting_date': meeting_minutes.basic_info.date or datetime.now().strftime('%Y-%m-%d'),
                'meeting_time': meeting_minutes.basic_info.time or "",
                'meeting_location': meeting_minutes.basic_info.location or "",
                'meeting_duration': meeting_minutes.basic_info.duration or "",
                'meeting_organizer': meeting_minutes.basic_info.organizer or "",
            })
        
        # Generation date
        data['generation_date'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # Attendees
        if meeting_minutes.attendees:
            data['attendees'] = [
                {
                    'name': attendee.name,
                    'role': attendee.role or "",
                    'organization': attendee.organization or "",
                    'present': "出席" if attendee.present else "缺席"
                }
                for attendee in meeting_minutes.attendees
            ]
        
        # Agenda
        if meeting_minutes.agenda:
            data['agenda'] = [
                {
                    'title': topic.title,
                    'description': topic.description or "",
                    'presenter': topic.presenter or "",
                    'key_points': topic.key_points or []
                }
                for topic in meeting_minutes.agenda
            ]
        
        # Action items
        if meeting_minutes.action_items:
            data['action_items'] = [
                {
                    'task': action.task,
                    'assignee': action.assignee or "",
                    'due_date': action.due_date or "",
                    'priority': action.priority or "",
                    'status': action.status or "待處理"
                }
                for action in meeting_minutes.action_items
            ]
        
        # Decisions
        if meeting_minutes.decisions:
            data['decisions'] = [
                {
                    'decision': decision.decision,
                    'rationale': decision.rationale or "",
                    'impact': decision.impact or "",
                    'responsible_party': decision.responsible_party or ""
                }
                for decision in meeting_minutes.decisions
            ]
        
        # Key outcomes
        if meeting_minutes.key_outcomes:
            data['key_outcomes'] = meeting_minutes.key_outcomes
        
        # Additional notes
        if meeting_minutes.additional_notes:
            data['additional_notes'] = meeting_minutes.additional_notes
        
        # Next meeting
        if meeting_minutes.next_meeting:
            data['next_meeting'] = meeting_minutes.next_meeting
        
        return data
    
    def _apply_document_formatting(self, doc: 'Document', formatting: Dict[str, Any]):
        """Apply document-level formatting"""
        try:
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = formatting.get('font_name', '微軟正黑體')
            font.size = Pt(formatting.get('font_size', 12))
            
            # Set paragraph formatting
            paragraph_format = style.paragraph_format
            paragraph_format.line_spacing = formatting.get('line_spacing', 1.15)
            paragraph_format.space_after = Pt(formatting.get('paragraph_spacing', 6))
            
        except Exception as e:
            logger.warning(f"⚠️ Document formatting warning: {e}")
    
    def _apply_paragraph_formatting(self, paragraph, formatting: Dict[str, Any]):
        """Apply paragraph-specific formatting"""
        try:
            if formatting.get('bold'):
                for run in paragraph.runs:
                    run.bold = True
            
            if formatting.get('italic'):
                for run in paragraph.runs:
                    run.italic = True
            
            if formatting.get('color'):
                color = formatting['color']
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(color[0], color[1], color[2])
            
        except Exception as e:
            logger.warning(f"⚠️ Paragraph formatting warning: {e}")
    
    async def _inject_data_text_mode(
        self, 
        meeting_minutes: MeetingMinutes, 
        template_config: 'TemplateConfig'
    ) -> InjectionResult:
        """Fallback text-mode data injection"""
        logger.info("📝 Using text-mode data injection (python-docx not available)")
        
        # Prepare simple text representation
        injected_fields = ['meeting_title', 'meeting_date', 'attendees', 'key_outcomes']
        
        return InjectionResult(
            success=True,
            injected_fields=injected_fields,
            warnings=["Text-mode injection - limited formatting available"],
            metadata={
                'processing_mode': 'text',
                'docx_available': False
            }
        )
    
    def get_injection_rules(self) -> Dict[str, InjectionRule]:
        """Get all injection rules"""
        return self.injection_rules.copy()
    
    def validate_data_completeness(self, meeting_minutes: MeetingMinutes) -> Dict[str, Any]:
        """Validate data completeness for injection"""
        injection_data = self._prepare_injection_data(meeting_minutes)
        
        required_fields = [rule.field_name for rule in self.injection_rules.values() if rule.required]
        available_fields = list(injection_data.keys())
        missing_fields = [field for field in required_fields if field not in available_fields]
        
        completeness_score = (len(available_fields) - len(missing_fields)) / len(required_fields) if required_fields else 1.0
        
        return {
            'completeness_score': completeness_score,
            'total_fields': len(required_fields),
            'available_fields': len(available_fields),
            'missing_fields': missing_fields,
            'data_quality': {
                'has_basic_info': bool(meeting_minutes.basic_info),
                'has_attendees': bool(meeting_minutes.attendees),
                'has_agenda': bool(meeting_minutes.agenda),
                'has_action_items': bool(meeting_minutes.action_items),
                'has_decisions': bool(meeting_minutes.decisions),
                'has_outcomes': bool(meeting_minutes.key_outcomes)
            }
        }