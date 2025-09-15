import logging
import tempfile
import os
import platform
import subprocess
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
try:
    from docx.oxml.parser import parse_xml
except ImportError:
    # Fallback for older versions
    from lxml import etree
    def parse_xml(xml_str):
        return etree.fromstring(xml_str)
from typing import Dict, Any, List, Tuple
from io import BytesIO
from datetime import datetime

logger = logging.getLogger(__name__)


class WordGenerator:
    def __init__(self):
        pass
    
    def generate_document(self, meeting_data: Dict[str, Any]) -> bytes:
        """Generate professional Word document from meeting data following SESE principles"""
        try:
            doc = Document()
            
            # Set document margins and formatting
            self._configure_document_layout(doc)
            
            # Add professional header with company branding area
            self._add_document_header(doc, meeting_data)
            
            # Add title with enhanced styling
            self._add_professional_title(doc, meeting_data.get("meeting_title", "會議記錄"))
            
            # Add executive summary (new)
            if meeting_data.get("summary"):
                self._add_executive_summary(doc, meeting_data["summary"])
            
            # Add comprehensive meeting info
            self._add_enhanced_meeting_info(doc, meeting_data)
            
            # Add attendees with enhanced formatting
            self._add_attendees(doc, meeting_data.get("attendees", []))
            
            # Add key topics with improved structure
            self._add_key_topics(doc, meeting_data.get("key_topics", []))
            
            # Add decisions with enhanced presentation
            self._add_decisions(doc, meeting_data.get("decisions", []))
            
            # Add action items with priority support
            self._add_enhanced_action_items(doc, meeting_data.get("action_items", []))
            
            # Add follow-up items (new)
            if meeting_data.get("follow_up_items"):
                self._add_follow_up_items(doc, meeting_data["follow_up_items"])
            
            # Add next meeting info with enhanced formatting
            if meeting_data.get("next_meeting"):
                self._add_next_meeting(doc, meeting_data["next_meeting"])
            
            # Add professional footer with metadata
            self._add_professional_footer(doc, meeting_data)
            
            # Save to bytes
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            logger.info("Professional Word document generated successfully")
            return doc_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Failed to generate Word document: {e}")
            raise
    
    def _configure_document_layout(self, doc: Document):
        """Configure professional document layout and styling"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.page_height = Inches(11.69)  # A4
            section.page_width = Inches(8.27)    # A4
    
    def _add_document_header(self, doc: Document, meeting_data: Dict[str, Any]):
        """Add professional document header with branding space"""
        header_table = doc.add_table(rows=1, cols=3)
        header_table.autofit = False
        
        # Left cell for company logo/name (placeholder)
        logo_cell = header_table.cell(0, 0)
        logo_p = logo_cell.paragraphs[0]
        logo_run = logo_p.add_run("Company Logo")  # Placeholder for logo
        logo_run.font.size = Pt(10)
        logo_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Center cell for document type
        center_cell = header_table.cell(0, 1)
        center_p = center_cell.paragraphs[0]
        center_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        center_run = center_p.add_run("Meeting Minutes")
        center_run.font.size = Pt(12)
        center_run.bold = True
        
        # Right cell for date/time
        date_cell = header_table.cell(0, 2)
        date_p = date_cell.paragraphs[0]
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        current_date = datetime.now().strftime("%Y-%m-%d")
        date_run = date_p.add_run(f"Generated: {current_date}")
        date_run.font.size = Pt(10)
        
        # Add line separator
        doc.add_paragraph("")
        separator = doc.add_paragraph()
        separator.add_run("─" * 80)
        doc.add_paragraph("")
    
    def _add_professional_title(self, doc: Document, title: str):
        """Add professionally styled title"""
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Enhanced title styling
        for run in title_paragraph.runs:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(31, 73, 125)  # Professional blue
            run.bold = True
        
        doc.add_paragraph("")
    
    def _add_executive_summary(self, doc: Document, summary: str):
        """Add executive summary section"""
        summary_heading = doc.add_heading("Executive Summary", level=1)
        self._style_section_heading(summary_heading)
        
        summary_p = doc.add_paragraph()
        summary_p.style.font.size = Pt(11)
        summary_p.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Add summary box styling
        summary_run = summary_p.add_run(summary)
        summary_run.font.size = Pt(11)
        
        doc.add_paragraph("")
    
    def _add_enhanced_meeting_info(self, doc: Document, meeting_data: Dict[str, Any]):
        """Add comprehensive meeting information with professional styling"""
        info_heading = doc.add_heading("Meeting Information", level=1)
        self._style_section_heading(info_heading)
        
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Enhanced meeting info data
        info_data = [
            ("Meeting Title", meeting_data.get("meeting_title", "General Meeting")),
            ("Meeting Type", meeting_data.get("meeting_type", "General").title()),
            ("Date", meeting_data.get("date", "Not specified")),
            ("Time", meeting_data.get("time", "Not specified")),
            ("Duration", meeting_data.get("duration", "Not specified")),
            ("Location", meeting_data.get("location", "Not specified"))
        ]
        
        for row_idx, (label, value) in enumerate(info_data):
            row = info_table.rows[row_idx]
            label_cell = row.cells[0]
            value_cell = row.cells[1]
            
            # Style label cell
            label_cell.text = label
            for paragraph in label_cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(31, 73, 125)
            
            # Style value cell
            value_cell.text = value or "Not specified"
            for paragraph in value_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
        
        # Apply table styling
        self._apply_table_styling(info_table)
        doc.add_paragraph("")
    
    def _add_enhanced_action_items(self, doc: Document, action_items: List[Dict[str, str]]):
        """Add action items with priority indicators and enhanced formatting"""
        action_heading = doc.add_heading("Action Items", level=1)
        self._style_section_heading(action_heading)
        
        if action_items:
            # Create enhanced table for action items
            action_table = doc.add_table(rows=1, cols=4)
            action_table.style = 'Table Grid'
            action_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # Header row with enhanced styling
            header_cells = action_table.rows[0].cells
            headers = ["Priority", "Task Description", "Assignee", "Deadline"]
            
            for i, header in enumerate(headers):
                cell = header_cells[i]
                cell.text = header
                # Style header cells
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(255, 255, 255)
                # Add background color to header
                self._set_cell_background_color(cell, RGBColor(31, 73, 125))
            
            # Add action items with priority indicators
            for action in action_items:
                row_cells = action_table.add_row().cells
                
                # Priority cell with color coding
                priority = action.get("priority", "medium").lower()
                priority_text = priority.upper()
                row_cells[0].text = priority_text
                
                # Color code priority
                priority_color = self._get_priority_color(priority)
                for paragraph in row_cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = priority_color
                
                # Task description
                row_cells[1].text = action.get("task", "")
                
                # Assignee
                row_cells[2].text = action.get("assignee", "Not assigned")
                
                # Deadline
                row_cells[3].text = action.get("deadline", "To be determined")
            
            self._apply_table_styling(action_table)
        else:
            no_actions = doc.add_paragraph("No action items identified in this meeting.")
            for run in no_actions.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph("")
    
    def _add_follow_up_items(self, doc: Document, follow_up_items: List[str]):
        """Add follow-up items section"""
        if not follow_up_items:
            return
            
        followup_heading = doc.add_heading("Follow-up Items", level=1)
        self._style_section_heading(followup_heading)
        
        for item in follow_up_items:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            run = p.add_run(item)
            run.font.size = Pt(11)
        
        doc.add_paragraph("")
    
    def _add_professional_footer(self, doc: Document, meeting_data: Dict[str, Any]):
        """Add comprehensive professional footer with metadata"""
        doc.add_paragraph("")
        doc.add_paragraph("─" * 80)
        
        # Document metadata table
        metadata_table = doc.add_table(rows=3, cols=2)
        
        current_time = datetime.now()
        generation_info = [
            ("Document Generated", current_time.strftime("%Y-%m-%d %H:%M:%S")),
            ("Generated By", "Rapid Minutes Export System v2.0"),
            ("Meeting Type", meeting_data.get("meeting_type", "General").title())
        ]
        
        for row_idx, (label, value) in enumerate(generation_info):
            row = metadata_table.rows[row_idx]
            label_cell = row.cells[0]
            value_cell = row.cells[1]
            
            label_cell.text = label
            value_cell.text = value
            
            # Style metadata
            for paragraph in label_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.bold = True
            
            for paragraph in value_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.italic = True
        
        # Add confidentiality notice
        doc.add_paragraph("")
        confidentiality = doc.add_paragraph()
        confidentiality.alignment = WD_ALIGN_PARAGRAPH.CENTER
        conf_run = confidentiality.add_run("This document contains confidential meeting information. Please handle with appropriate care.")
        conf_run.font.size = Pt(8)
        conf_run.italic = True
        conf_run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _style_section_heading(self, heading):
        """Apply consistent styling to section headings"""
        for run in heading.runs:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(31, 73, 125)
            run.bold = True
    
    def _apply_table_styling(self, table):
        """Apply professional styling to tables"""
        table.autofit = True
        
        # Set consistent row height and cell padding
        for row in table.rows:
            row.height = Inches(0.3)
            for cell in row.cells:
                # Add cell padding
                cell.margin_top = Inches(0.05)
                cell.margin_bottom = Inches(0.05)
                cell.margin_left = Inches(0.1)
                cell.margin_right = Inches(0.1)
    
    def _get_priority_color(self, priority: str) -> RGBColor:
        """Get color coding for priority levels"""
        priority_colors = {
            "high": RGBColor(220, 53, 69),    # Red
            "medium": RGBColor(255, 193, 7),  # Yellow/Orange
            "low": RGBColor(40, 167, 69)      # Green
        }
        return priority_colors.get(priority.lower(), RGBColor(128, 128, 128))
    
    def _set_cell_background_color(self, cell, color: RGBColor):
        """Set background color for table cell"""
        try:
            # Create shading element
            shading_element = parse_xml(r'<w:shd {} w:fill="{:02x}{:02x}{:02x}"/>'.format(
                nsdecls('w'), color.r, color.g, color.b))
            cell._tc.get_or_add_tcPr().append(shading_element)
        except Exception:
            # Fallback if advanced styling fails
            pass
    
    def _add_meeting_info(self, doc: Document, meeting_data: Dict[str, Any]):
        """Add basic meeting information"""
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        # Meeting info data
        info_data = [
            ("會議主題", meeting_data.get("meeting_title", "一般會議")),
            ("會議日期", meeting_data.get("date", "")),
            ("會議時間", meeting_data.get("time", "")),
            ("會議地點", meeting_data.get("location", ""))
        ]
        
        for row_idx, (label, value) in enumerate(info_data):
            row = info_table.rows[row_idx]
            label_cell = row.cells[0]
            value_cell = row.cells[1]
            
            label_cell.text = label
            value_cell.text = value or "未指定"
            
            # Format label cell
            label_run = label_cell.paragraphs[0].runs[0]
            label_run.bold = True
            
        doc.add_paragraph("")  # Add spacing
    
    def _add_attendees(self, doc: Document, attendees: List[str]):
        """Add attendees section"""
        if not attendees:
            return
            
        doc.add_heading("出席人員", level=1)
        
        if attendees:
            for attendee in attendees:
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(attendee)
        else:
            doc.add_paragraph("無記錄")
        
        doc.add_paragraph("")  # Add spacing
    
    def _add_key_topics(self, doc: Document, key_topics: List[str]):
        """Add key discussion topics"""
        doc.add_heading("討論議題", level=1)
        
        if key_topics:
            for idx, topic in enumerate(key_topics, 1):
                p = doc.add_paragraph()
                p.style = 'List Number'
                p.add_run(topic)
        else:
            doc.add_paragraph("無記錄的討論議題")
        
        doc.add_paragraph("")  # Add spacing
    
    def _add_decisions(self, doc: Document, decisions: List[str]):
        """Add decisions made"""
        doc.add_heading("決議事項", level=1)
        
        if decisions:
            for decision in decisions:
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                # Add bullet point manually if style doesn't work
                p.add_run(f"• {decision}")
        else:
            doc.add_paragraph("本次會議無明確決議事項")
        
        doc.add_paragraph("")  # Add spacing
    
    def _add_action_items(self, doc: Document, action_items: List[Dict[str, str]]):
        """Add action items"""
        doc.add_heading("行動事項", level=1)
        
        if action_items:
            # Create table for action items
            action_table = doc.add_table(rows=1, cols=3)
            action_table.style = 'Table Grid'
            
            # Header row
            header_cells = action_table.rows[0].cells
            header_cells[0].text = "事項"
            header_cells[1].text = "負責人"
            header_cells[2].text = "期限"
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add action items
            for action in action_items:
                row_cells = action_table.add_row().cells
                row_cells[0].text = action.get("task", "")
                row_cells[1].text = action.get("assignee", "未指定")
                row_cells[2].text = action.get("deadline", "待確認")
        else:
            doc.add_paragraph("無行動事項")
        
        doc.add_paragraph("")  # Add spacing
    
    def _add_next_meeting(self, doc: Document, next_meeting: str):
        """Add next meeting information"""
        doc.add_heading("下次會議", level=1)
        doc.add_paragraph(next_meeting)
        doc.add_paragraph("")  # Add spacing
    
    def _add_footer(self, doc: Document):
        """Add document footer"""
        doc.add_paragraph("")
        doc.add_paragraph("─" * 50)
        
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add generation info
        current_time = datetime.now().strftime("%Y年%m月%d日 %H:%M")
        footer_p.add_run(f"本記錄由 Rapid Minutes Export 系統生成於 {current_time}")
        
        # Make footer text smaller and italic
        for run in footer_p.runs:
            run.font.size = Pt(9)
            run.italic = True
    
    def create_template(self) -> bytes:
        """Create a template Word document"""
        try:
            doc = Document()
            
            # Add title placeholder
            title = doc.add_heading("會議記錄範本", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add template sections
            doc.add_heading("會議基本資訊", level=1)
            doc.add_paragraph("會議主題：[待填入]")
            doc.add_paragraph("會議日期：[待填入]")
            doc.add_paragraph("會議時間：[待填入]")
            doc.add_paragraph("會議地點：[待填入]")
            
            doc.add_heading("出席人員", level=1)
            doc.add_paragraph("• [待填入]")
            
            doc.add_heading("討論議題", level=1)
            doc.add_paragraph("1. [待填入]")
            
            doc.add_heading("決議事項", level=1)
            doc.add_paragraph("• [待填入]")
            
            doc.add_heading("行動事項", level=1)
            action_table = doc.add_table(rows=2, cols=3)
            action_table.style = 'Table Grid'
            
            header_cells = action_table.rows[0].cells
            header_cells[0].text = "事項"
            header_cells[1].text = "負責人"
            header_cells[2].text = "期限"
            
            sample_cells = action_table.rows[1].cells
            sample_cells[0].text = "[待填入]"
            sample_cells[1].text = "[待填入]"
            sample_cells[2].text = "[待填入]"
            
            # Save template
            template_buffer = BytesIO()
            doc.save(template_buffer)
            template_buffer.seek(0)
            
            return template_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Failed to create template: {e}")
            raise
    
    def generate_pdf_from_word(self, word_content: bytes) -> bytes:
        """Convert Word document to PDF using multiple methods for maximum compatibility"""
        try:
            # Try different PDF conversion methods based on system capabilities
            conversion_methods = [
                self._convert_with_docx2pdf,
                self._convert_with_libreoffice,
                self._convert_with_pandoc
            ]
            
            for method in conversion_methods:
                try:
                    pdf_content = method(word_content)
                    if pdf_content:
                        logger.info(f"PDF conversion successful using {method.__name__}")
                        return pdf_content
                except Exception as e:
                    logger.warning(f"PDF conversion failed with {method.__name__}: {e}")
                    continue
            
            # If all methods fail, raise an exception
            raise Exception("All PDF conversion methods failed")
            
        except Exception as e:
            logger.error(f"Failed to convert Word to PDF: {e}")
            raise
    
    def _convert_with_docx2pdf(self, word_content: bytes) -> bytes:
        """Convert using docx2pdf library (Windows optimized)"""
        try:
            from docx2pdf import convert
            
            with tempfile.TemporaryDirectory() as temp_dir:
                # Save Word content to temporary file
                word_path = os.path.join(temp_dir, "document.docx")
                with open(word_path, "wb") as f:
                    f.write(word_content)
                
                # Convert to PDF
                pdf_path = os.path.join(temp_dir, "document.pdf")
                convert(word_path, pdf_path)
                
                # Read PDF content
                with open(pdf_path, "rb") as f:
                    return f.read()
                    
        except ImportError:
            logger.warning("docx2pdf library not available")
            raise Exception("docx2pdf not installed")
        except Exception as e:
            logger.error(f"docx2pdf conversion failed: {e}")
            raise
    
    def _convert_with_libreoffice(self, word_content: bytes) -> bytes:
        """Convert using LibreOffice command line (Linux/Mac optimized)"""
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                # Save Word content to temporary file
                word_path = os.path.join(temp_dir, "document.docx")
                with open(word_path, "wb") as f:
                    f.write(word_content)
                
                # Use LibreOffice to convert
                cmd = [
                    "libreoffice", "--headless", "--convert-to", "pdf",
                    "--outdir", temp_dir, word_path
                ]
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                
                if result.returncode != 0:
                    raise Exception(f"LibreOffice conversion failed: {result.stderr}")
                
                # Read PDF content
                pdf_path = os.path.join(temp_dir, "document.pdf")
                if not os.path.exists(pdf_path):
                    raise Exception("PDF file was not created")
                
                with open(pdf_path, "rb") as f:
                    return f.read()
                    
        except subprocess.TimeoutExpired:
            raise Exception("LibreOffice conversion timed out")
        except FileNotFoundError:
            raise Exception("LibreOffice not found in system PATH")
        except Exception as e:
            logger.error(f"LibreOffice conversion failed: {e}")
            raise
    
    def _convert_with_pandoc(self, word_content: bytes) -> bytes:
        """Convert using pandoc (universal fallback)"""
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                # Save Word content to temporary file
                word_path = os.path.join(temp_dir, "document.docx")
                with open(word_path, "wb") as f:
                    f.write(word_content)
                
                # Use pandoc to convert
                pdf_path = os.path.join(temp_dir, "document.pdf")
                cmd = [
                    "pandoc", word_path, "-o", pdf_path,
                    "--pdf-engine=xelatex"  # or pdflatex
                ]
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=45)
                
                if result.returncode != 0:
                    raise Exception(f"Pandoc conversion failed: {result.stderr}")
                
                # Read PDF content
                if not os.path.exists(pdf_path):
                    raise Exception("PDF file was not created by pandoc")
                
                with open(pdf_path, "rb") as f:
                    return f.read()
                    
        except subprocess.TimeoutExpired:
            raise Exception("Pandoc conversion timed out")
        except FileNotFoundError:
            raise Exception("Pandoc not found in system PATH")
        except Exception as e:
            logger.error(f"Pandoc conversion failed: {e}")
            raise
    
    def generate_documents_bundle(self, meeting_data: Dict[str, Any]) -> Tuple[bytes, bytes]:
        """Generate both Word and PDF documents in one operation - implementing SESE efficiency principle"""
        try:
            # Generate Word document
            logger.info("Generating Word document...")
            word_content = self.generate_document(meeting_data)
            
            # Generate PDF from Word content
            logger.info("Converting to PDF...")
            pdf_content = self.generate_pdf_from_word(word_content)
            
            logger.info("Successfully generated both Word and PDF documents")
            return word_content, pdf_content
            
        except Exception as e:
            logger.error(f"Failed to generate document bundle: {e}")
            # Return Word document even if PDF fails
            try:
                word_content = self.generate_document(meeting_data)
                logger.warning("PDF generation failed, returning Word document only")
                return word_content, b''
            except Exception:
                logger.error("Both Word and PDF generation failed")
                raise
    
    def check_pdf_conversion_capabilities(self) -> Dict[str, bool]:
        """Check which PDF conversion methods are available on the system"""
        capabilities = {
            "docx2pdf": False,
            "libreoffice": False,
            "pandoc": False,
            "system_info": {
                "platform": platform.system(),
                "python_version": platform.python_version()
            }
        }
        
        # Check docx2pdf
        try:
            import docx2pdf
            capabilities["docx2pdf"] = True
        except ImportError:
            pass
        
        # Check LibreOffice
        try:
            result = subprocess.run(["libreoffice", "--version"], 
                                  capture_output=True, timeout=5)
            capabilities["libreoffice"] = result.returncode == 0
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        
        # Check Pandoc
        try:
            result = subprocess.run(["pandoc", "--version"], 
                                  capture_output=True, timeout=5)
            capabilities["pandoc"] = result.returncode == 0
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        
        logger.info(f"PDF conversion capabilities: {capabilities}")
        return capabilities
    
    def validate_document_output(self, document_content: bytes, format_type: str) -> Dict[str, Any]:
        """Validate generated document according to ICE principle - comprehensive coverage"""
        validation_result = {
            "is_valid": False,
            "file_size": len(document_content),
            "format": format_type,
            "issues": [],
            "quality_score": 0
        }
        
        try:
            # Basic size validation
            if len(document_content) == 0:
                validation_result["issues"].append("Document is empty")
                return validation_result
            
            # Minimum size check (reasonable documents should be at least 1KB)
            if len(document_content) < 1024:
                validation_result["issues"].append("Document appears too small")
                validation_result["quality_score"] -= 20
            
            # Format-specific validation
            if format_type.lower() == "docx":
                validation_result.update(self._validate_word_document(document_content))
            elif format_type.lower() == "pdf":
                validation_result.update(self._validate_pdf_document(document_content))
            
            # Calculate overall quality score
            base_score = 100
            base_score -= len(validation_result["issues"]) * 10
            validation_result["quality_score"] = max(0, min(100, base_score))
            validation_result["is_valid"] = validation_result["quality_score"] >= 70
            
            return validation_result
            
        except Exception as e:
            validation_result["issues"].append(f"Validation error: {str(e)}")
            return validation_result
    
    def _validate_word_document(self, word_content: bytes) -> Dict[str, Any]:
        """Validate Word document structure"""
        result = {"format_specific_issues": []}
        
        try:
            # Try to open the document to validate structure
            with tempfile.TemporaryDirectory() as temp_dir:
                word_path = os.path.join(temp_dir, "test.docx")
                with open(word_path, "wb") as f:
                    f.write(word_content)
                
                # Load and validate document
                doc = Document(word_path)
                
                # Check for content
                if len(doc.paragraphs) == 0:
                    result["format_specific_issues"].append("No paragraphs found")
                
                # Check for tables (action items should create tables)
                table_count = len(doc.tables)
                if table_count == 0:
                    result["format_specific_issues"].append("No tables found - may lack structured content")
                
                # Check document structure
                heading_count = sum(1 for p in doc.paragraphs if p.style.name.startswith('Heading'))
                if heading_count < 3:
                    result["format_specific_issues"].append("Few headings found - document may lack structure")
                
        except Exception as e:
            result["format_specific_issues"].append(f"Word document validation failed: {str(e)}")
        
        return result
    
    def _validate_pdf_document(self, pdf_content: bytes) -> Dict[str, Any]:
        """Validate PDF document structure"""
        result = {"format_specific_issues": []}
        
        # Basic PDF header validation
        if not pdf_content.startswith(b'%PDF-'):
            result["format_specific_issues"].append("Invalid PDF header")
        
        # Check for PDF trailer
        if b'%%EOF' not in pdf_content[-100:]:
            result["format_specific_issues"].append("PDF trailer not found - file may be corrupted")
        
        # Size reasonableness check
        if len(pdf_content) > 10 * 1024 * 1024:  # 10MB
            result["format_specific_issues"].append("PDF file unusually large")
        
        return result